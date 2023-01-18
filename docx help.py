import docx
import os.path

import config as config
from configobj import ConfigObj
import configparser
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.shared import Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
import docx

config = configparser.ConfigParser()
config.read('D:\FTP_Home\LocalUser\S_4_1\KH_PostProcessing\KH_Process1\config.ini')
filepath = config['default']['path1']
filename = config['default']['filename']
#config = ConfigObj('config.ini')
#filepath = 'D:\\FTP_Home\\LocalUser\\s1000d4_0\\channels\\output\\word\\docx\\'
#filename = 'Book_test_two.docx'

#file1 = open("D:\FTP_Home\LocalUser\s1000d4_0\KH_PostProcessing\KH_Process1\\test.txt","r")
#lines = file1.readlines()



#filepath = ( "".join([l.rstrip() for l in lines[0]]) )
#filename = ( "".join([l.rstrip() for l in lines[1]]) )


#print(filepath)
#print(filename)

#file1.close()


fullname = filepath + filename

#print(fullname)

document = Document(fullname)
paragraphs = document.paragraphs

document = Document(fullname)
paragraphs = document.paragraphs


paragraphs[0].text = ""
paragraphs[1].text = ""
paragraphs[2].text = ""

paragraphs[0].add_run('1')
paragraphs[1].add_run('Title')
paragraphs[2].add_run('1.Start')
paragraphs[3].text = "_"

paragraphs[0].style = "PAGE"
paragraphs[1].style = "PARA"
paragraphs[2].style = "START"
paragraphs[3].style = "PARA"

#paragraphs[3].add_run('Title1')

#paragraphs[4].add_run('1. Start ')
#paragraphs[3].style = "Title2"
#paragraphs[4].style = "START"


#filepath = 'Output/' + filename +'-01.docx'

filepath = filepath  + filename.rsplit( ".", 1 )[ 0 ] +'-01.docx'
document.save(filepath)
print("1")


